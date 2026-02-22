import os
import glob
import sys

try:
    from docx import Document
except Exception:
    print("python-docx not installed. Please run: pip install python-docx")
    sys.exit(2)

WORKSPACE = r"C:\Users\Serenay\Desktop\TEZ"
TEZ_DIR = os.path.join(WORKSPACE, "TEZ")
OUTPUT_PATH = os.path.join(WORKSPACE, "tez_merged.docx")

def md_heading_level(line: str):
    # Count leading '#' characters to determine heading level
    stripped = line.lstrip()
    if not stripped.startswith("#"):
        return 0, line
    hashes = 0
    for ch in stripped:
        if ch == "#":
            hashes += 1
        else:
            break
    # extract text after hashes and optional space
    text = stripped[hashes:].strip()
    return hashes, text

def main():
    md_files = sorted(glob.glob(os.path.join(TEZ_DIR, "*.md")))
    if not md_files:
        print("No markdown files found in", TEZ_DIR)
        sys.exit(1)

    doc = Document()

    for md_path in md_files:
        base = os.path.basename(md_path)
        title = base.replace(".md", "")
        doc.add_heading(title, level=1)

        with open(md_path, "r", encoding="utf-8") as fh:
            for raw in fh:
                line = raw.rstrip("\n")
                if not line.strip():
                    # preserve blank lines
                    doc.add_paragraph("")
                    continue

                level, text = md_heading_level(line)
                if level > 0:
                    # map Markdown heading levels to Word heading levels (1-4)
                    word_level = min(max(level, 1), 4)
                    doc.add_heading(text, level=word_level)
                else:
                    # simple paragraph for non-heading lines
                    doc.add_paragraph(line)

        # add a page break between files
        doc.add_page_break()

    doc.save(OUTPUT_PATH)
    print("Merged and saved:", OUTPUT_PATH)

if __name__ == "__main__":
    main()

